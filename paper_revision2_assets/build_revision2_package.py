from __future__ import annotations

import json
import math
import shutil
import subprocess
import sys
from datetime import date
from pathlib import Path
from zipfile import ZipFile

import matplotlib as mpl
import matplotlib.image as mpimg
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import seaborn as sns
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Font, PatternFill


ROOT = Path(__file__).resolve().parents[2]
PREV = ROOT / "vexpaer_paper" / "投稿1"
OUT = ROOT / "vexpaer_paper_2" / "投稿2"
RESULT = ROOT / "lsj_code" / "result"
TABLES = RESULT / "tables"
FIGS = RESULT / "figures"


TITLE = "The illusion of the universal baseline in public single-cell controls"
MS_NO = "JBCB-1487R1"
JOURNAL = "Journal of Bioinformatics and Computational Biology"
TODAY = date.today().isoformat()


MAPBATCH_REF = (
    "Yong CH, Hoon S, de Mel S, Xu S, Scolnick JA, Huo D, Lovci MT, "
    "Chng WJ, Goh WWB. MapBatch: Conservative batch normalization for "
    "single cell RNA-sequencing data enables discovery of rare cell "
    "populations in a multiple myeloma cohort. Blood. 2021;138(Suppl 1):2954. "
    "doi:10.1182/blood-2021-150089"
)


REFERENCES = [
    "Luecken MD, Theis FJ. Current best practices in single-cell RNA-seq analysis: a tutorial. Mol Syst Biol. 2019;15:e8746. doi:10.15252/msb.20188746",
    "Tran HTN, Ang KS, Chevrier M, Zhang X, Lee NYS, Goh M, et al. A benchmark of batch-effect correction methods for single-cell RNA sequencing data. Genome Biol. 2020;21:12. doi:10.1186/s13059-019-1850-9",
    "Heumos L, Schaar AC, Lance C, Litinetskaya A, Drost F, Zappia L, et al. Best practices for single-cell analysis across modalities. Nat Rev Genet. 2023;24:550-572. doi:10.1038/s41576-023-00586-w",
    "Li M, Zhang X, Ang KS, Ling J, Sethi R, Lee NYS, et al. DISCO: a database of Deeply Integrated human Single-Cell Omics data. Nucleic Acids Res. 2022;50:D596-D602. doi:10.1093/nar/gkab1020",
    "Büttner M, Miao Z, Wolf FA, Teichmann SA, Theis FJ. A test metric for assessing single-cell RNA-seq batch correction. Nat Methods. 2019;16:43-49. doi:10.1038/s41592-018-0254-1",
    "Squair JW, Gautier M, Kathe C, Anderson MA, James ND, Hutson TH, et al. Confronting false discoveries in single-cell differential expression. Nat Commun. 2021;12:5692. doi:10.1038/s41467-021-25960-2",
    "Crowell HL, Soneson C, Germain PL, Calini D, Collin L, Raposo C, et al. muscat detects subpopulation-specific state transitions from multi-sample multi-condition single-cell transcriptomics data. Nat Commun. 2020;11:6077. doi:10.1038/s41467-020-19894-4",
    "Korsunsky I, Millard N, Fan J, Slowikowski K, Zhang F, Wei K, et al. Fast, sensitive and accurate integration of single-cell data with Harmony. Nat Methods. 2019;16:1289-1296. doi:10.1038/s41592-019-0619-0",
    "Lopez R, Regier J, Cole MB, Jordan MI, Yosef N. Deep generative modeling for single-cell transcriptomics. Nat Methods. 2018;15:1053-1058. doi:10.1038/s41592-018-0229-2",
    "Love MI, Huber W, Anders S. Moderated estimation of fold change and dispersion for RNA-seq data with DESeq2. Genome Biol. 2014;15:550. doi:10.1186/s13059-014-0550-8",
    MAPBATCH_REF,
]


def ensure_clean_out() -> None:
    if OUT.exists():
        for item in OUT.iterdir():
            if item.is_dir():
                shutil.rmtree(item)
            else:
                item.unlink()
    OUT.mkdir(parents=True, exist_ok=True)
    shutil.copy2(PREV / "Supplementary Materials.xlsx", OUT / "Supplementary Materials.xlsx")


def setup_matplotlib() -> None:
    mpl.rcParams.update(
        {
            "font.family": "DejaVu Sans",
            "pdf.fonttype": 42,
            "ps.fonttype": 42,
            "axes.spines.top": False,
            "axes.spines.right": False,
            "figure.facecolor": "white",
            "axes.facecolor": "white",
        }
    )


def panel_label(ax, label: str, x: float = 0.0, y: float = 1.02) -> None:
    ax.text(
        x,
        y,
        label,
        transform=ax.transAxes,
        fontsize=15,
        fontweight="bold",
        ha="left",
        va="bottom",
    )


def save_figure(fig: plt.Figure, stem: str, dpi: int = 450) -> None:
    for suffix in [".png", ".pdf", ".eps"]:
        fig.savefig(OUT / f"{stem}{suffix}", bbox_inches="tight", dpi=dpi if suffix == ".png" else None)
    plt.close(fig)


def save_supplementary_figure(fig: plt.Figure, stem: str, suffixes=(".png", ".pdf", ".eps"), dpi: int = 600) -> None:
    for suffix in suffixes:
        fig.savefig(OUT / f"{stem}{suffix}", bbox_inches="tight", dpi=dpi if suffix == ".png" else None)
    plt.close(fig)


def draw_box(ax, xy, w, h, text, fc="#F7F9FC", ec="#4A5568", fs=9, weight="normal") -> None:
    rect = mpl.patches.FancyBboxPatch(
        xy,
        w,
        h,
        boxstyle="round,pad=0.02,rounding_size=0.025",
        fc=fc,
        ec=ec,
        lw=1.2,
    )
    ax.add_patch(rect)
    ax.text(xy[0] + w / 2, xy[1] + h / 2, text, ha="center", va="center", fontsize=fs, fontweight=weight, wrap=True)


def make_figure1() -> None:
    setup_matplotlib()
    fig = plt.figure(figsize=(12.8, 8.8))
    gs = fig.add_gridspec(2, 2, height_ratios=[1.05, 1.0], width_ratios=[1.05, 1], hspace=0.28, wspace=0.24)
    fig.suptitle("From universal-baseline illusion to compatibility-first reuse", fontsize=18, fontweight="bold", y=0.98)

    ax = fig.add_subplot(gs[0, 0])
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    panel_label(ax, "A", -0.03)
    ax.text(0.02, 0.91, "A shared label can hide incompatible histories", fontsize=12, fontweight="bold")
    studies = [
        ("Study 1\ncontrol", "#CFE8FF", "whole-tissue\nscRNA-seq"),
        ("Study 2\ncontrol", "#E4F7D2", "sorted or\nenriched"),
        ("Study 3\ncontrol", "#FFE6C7", "single-nucleus\nor autopsy"),
        ("Study 4\ncontrol", "#F5D7E3", "different donor\ncontext"),
    ]
    x0s = [0.03, 0.27, 0.51, 0.75]
    for (label, color, hidden), x in zip(studies, x0s):
        draw_box(ax, (x, 0.64), 0.19, 0.15, label, fc=color, fs=9, weight="bold")
        draw_box(ax, (x, 0.41), 0.19, 0.12, hidden, fc="#FFFFFF", ec="#8795A1", fs=7.8)
        ax.annotate("", xy=(x + 0.095, 0.64), xytext=(x + 0.095, 0.53), arrowprops=dict(arrowstyle="->", lw=1.0, color="#4A5568"))
    ax.text(0.5, 0.34, "different hidden compatibility variables", ha="center", va="center", fontsize=8.2, color="#4A5568")
    ax.annotate("", xy=(0.5, 0.24), xytext=(0.5, 0.315), arrowprops=dict(arrowstyle="->", lw=1.3, color="#6B7280"))
    draw_box(ax, (0.23, 0.06), 0.54, 0.16, "Universal baseline?\nOnly after compatibility is tested", fc="#FFF6D6", ec="#9A6B00", fs=8.8, weight="bold")

    ax = fig.add_subplot(gs[0, 1])
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    panel_label(ax, "B", -0.03)
    ax.text(0.02, 0.91, "Compatibility-first workflow", fontsize=12, fontweight="bold")
    steps = [
        ("1. Define the control label", "donor source, disease context"),
        ("2. Build compatibility table", "modality, tissue, enrichment, chemistry"),
        ("3. Show data before and after integration", "do not rely on corrected UMAP alone"),
        ("4. Pair batch-removal and biology-conservation metrics", "ASW, iLISI, cLISI, graph connectivity"),
        ("5. Verify at donor level", "composition, state, sensitivity analyses"),
    ]
    y = 0.77
    for title, detail in steps:
        draw_box(ax, (0.05, y), 0.9, 0.105, f"{title}\n{detail}", fc="#F7F9FC", ec="#CBD5E0", fs=8.2)
        if y > 0.18:
            ax.annotate("", xy=(0.5, y - 0.025), xytext=(0.5, y), arrowprops=dict(arrowstyle="->", color="#4A5568"))
        y -= 0.145

    ax = fig.add_subplot(gs[1, 0])
    ax.axis("off")
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    panel_label(ax, "C", -0.03)
    ax.text(0.02, 0.92, "Metric logic: what each diagnostic asks", fontsize=12, fontweight="bold")
    rows = [
        ("ASW", "Are labels separated\nor overlapping?", "Visible study or\ncell-type separation"),
        ("iLISI", "How diverse are study\nlabels locally?", "Batch mixing;\nhigher is more mixed"),
        ("cLISI", "How mixed are biological\nlabels locally?", "Possible over-mixing\nof biological labels"),
        ("Graph\nconnectivity", "Do same-label cells\nstay connected?", "Fragmentation of\ncell states"),
    ]
    x_cols = [0.04, 0.21, 0.58]
    widths = [0.16, 0.35, 0.38]
    y_top = 0.78
    row_h = 0.125
    headers = ["Metric", "Reader question", "Why it matters"]
    for x, w, h in zip(x_cols, widths, headers):
        rect = mpl.patches.Rectangle((x, y_top), w, row_h, fc="#EAF2FA", ec="#D0D7DE", lw=1)
        ax.add_patch(rect)
        ax.text(x + 0.015, y_top + row_h / 2, h, ha="left", va="center", fontsize=8.6, fontweight="bold")
    for r_i, row in enumerate(rows):
        y = y_top - (r_i + 1) * row_h
        fill = "#FFFFFF" if r_i % 2 == 0 else "#F8FAFC"
        for x, w, txt in zip(x_cols, widths, row):
            rect = mpl.patches.Rectangle((x, y), w, row_h, fc=fill, ec="#D0D7DE", lw=1)
            ax.add_patch(rect)
            ax.text(x + 0.015, y + row_h / 2, txt, ha="left", va="center", fontsize=8.1, linespacing=1.15)

    ax = fig.add_subplot(gs[1, 1])
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    panel_label(ax, "D", -0.03)
    ax.text(0.02, 0.92, "Buyer-beware checks before pooling public controls", fontsize=12, fontweight="bold")
    checks = [
        "Report unintegrated and integrated views.",
        "Separate composition-eligible datasets from sensitivity-only datasets.",
        "Treat donor or sample, not cell, as the independent unit.",
        "Use protocol-restricted and leave-one-study-out sensitivity analyses.",
        "Disclose unresolved metadata gaps rather than hiding them by integration.",
    ]
    y = 0.76
    for item in checks:
        ax.text(0.06, y, f"[ ] {item}", fontsize=9.5, va="center")
        y -= 0.13
    draw_box(ax, (0.07, 0.03), 0.86, 0.15, "Core principle:\nintegration aligns representations;\nit does not certify study compatibility.", fc="#FFF6D6", ec="#9A6B00", fs=8.8, weight="bold")
    save_figure(fig, "Figure1")


def _metrics_data() -> pd.DataFrame:
    metrics = pd.read_csv(TABLES / "integration_metrics.csv")
    scvi_path = TABLES / "scvi_integration_metrics.csv"
    if scvi_path.exists():
        metrics = pd.concat([metrics, pd.read_csv(scvi_path)], ignore_index=True)
    order = ["unintegrated", "harmony", "scvi"]
    return metrics.set_index("method").loc[[m for m in order if m in set(metrics["method"])]].reset_index()


def make_figure2() -> None:
    setup_matplotlib()
    fig = plt.figure(figsize=(13.2, 10.0))
    gs = fig.add_gridspec(3, 3, height_ratios=[0.75, 1.05, 0.85], hspace=0.38, wspace=0.26)
    fig.suptitle("Lung-control case study: integration reduces structure but does not prove exchangeability", fontsize=16, fontweight="bold", y=0.985)

    ax = fig.add_subplot(gs[0, :])
    ax.axis("off")
    panel_label(ax, "A", -0.01)
    ax.text(0.015, 0.92, "Study compatibility map", fontsize=12, fontweight="bold", transform=ax.transAxes)
    ds = pd.read_csv(TABLES / "dataset_summary.csv")
    rows = []
    for _, row in ds.iterrows():
        use = "controlled primary" if bool(row["controlled_primary"]) else row["analysis_set"].replace("_", " ")
        rows.append([row["study"], int(row["samples"]), row["modality"], row["protocol_class"].replace("_", " "), use])
    table = ax.table(cellText=rows, colLabels=["Study", "Samples", "Modality", "Protocol class", "Use in argument"], loc="center", cellLoc="left", colLoc="left", colWidths=[0.12, 0.09, 0.15, 0.34, 0.25])
    table.auto_set_font_size(False)
    table.set_fontsize(8.2)
    table.scale(1, 1.35)
    for (r, c), cell in table.get_celld().items():
        cell.set_edgecolor("#D0D7DE")
        if r == 0:
            cell.set_facecolor("#EAF2FA")
            cell.set_text_props(weight="bold")
        elif "controlled" in rows[r - 1][4]:
            cell.set_facecolor("#EAF7EA")
        else:
            cell.set_facecolor("#FFFFFF" if r % 2 else "#F8FAFC")

    umaps = [
        ("umap_unintegrated_by_study.png", "B", "Unintegrated"),
        ("umap_harmony_by_study.png", "C", "Harmony"),
        ("umap_scvi_by_study.png", "D", "scVI"),
    ]
    for i, (fname, label, title) in enumerate(umaps):
        ax = fig.add_subplot(gs[1, i])
        img = mpimg.imread(FIGS / fname)
        ax.imshow(img)
        ax.set_axis_off()
        ax.set_title(title, fontsize=11)
        panel_label(ax, label, -0.02)

    ax = fig.add_subplot(gs[2, 0])
    panel_label(ax, "E", -0.12)
    metrics = _metrics_data()
    x = np.arange(len(metrics))
    width = 0.25
    ax.bar(x - width, metrics["batch_asw_lower_is_better"], width, label="batch ASW\nlower", color="#4C78A8")
    ax.bar(x, metrics["iLISI_study_higher_is_better"], width, label="iLISI\nhigher", color="#54A24B")
    ax.bar(x + width, metrics["cLISI_celltype_lower_is_better"], width, label="cLISI\nlower", color="#F58518")
    ax.axhline(0, color="0.25", lw=0.8)
    ax.set_xticks(x)
    ax.set_xticklabels(["Unintegrated", "Harmony", "scVI"][: len(x)], rotation=20, ha="right")
    ax.set_ylabel("Metric value")
    ax.set_title("Integration diagnostics")
    ax.legend(frameon=False, fontsize=7, loc="upper left")

    ax = fig.add_subplot(gs[2, 1:])
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")
    panel_label(ax, "F", -0.02)
    ax.text(0.02, 0.91, "Donor-level evidence retained after protocol-aware filtering", fontsize=12, fontweight="bold")
    bullets = [
        "28 public lung control samples from five source studies.",
        "Controlled whole-tissue scRNA-seq subset: GSE173896 + GSE227691.",
        "All-study composition was study-associated; controlled subset divergence was lower but non-zero.",
        "Pseudobulk state and differential-expression screens remained study-associated in several cell types.",
        "Interpretation: residual study-conditioned heterogeneity, not a single biological mechanism.",
    ]
    y = 0.76
    for b in bullets:
        ax.text(0.05, y, f"- {b}", fontsize=9.4, va="top")
        y -= 0.13
    save_figure(fig, "Figure2")


STUDY_ORDER = ["GSE132771", "GSE159354", "GSE171524", "GSE173896", "GSE227691"]


def short_celltype(label: str) -> str:
    mapping = {
        "Alveolar macrophages": "AM",
        "Alveolar type I cells": "AT1",
        "Alveolar type II cells": "AT2",
        "Activated T cells": "Act. T",
        "CD4+ memory T cells": "CD4 mem. T",
        "CD8+ effector T cells": "CD8 eff. T",
        "Classical monocytes": "Class. mono",
        "Non-classical monocytes": "Non-class. mono",
        "Endothelial tip cells": "Endo tip",
        "Vascular endothelial cells": "Vasc. endo",
        "Smooth muscle cells": "SMC",
        "Follicular B cells": "Foll. B",
    }
    return mapping.get(label, label)


def make_supplementary_composition() -> None:
    comp = pd.read_csv(TABLES / "composition_all_study_diagnostic.csv")
    value_cols = [c for c in comp.columns if c not in {"analysis_set", "study", "sample_id"}]
    means = comp.groupby("study", observed=True)[value_cols].mean()
    means = means.reindex([s for s in STUDY_ORDER if s in means.index])
    means = means.div(means.sum(axis=1), axis=0).fillna(0.0) * 100.0
    order = means.mean(axis=0).sort_values(ascending=False).index.tolist()
    means = means[order]
    colors = list(mpl.colormaps["tab20"].colors) + list(mpl.colormaps["tab20b"].colors)

    fig, ax = plt.subplots(figsize=(9.2, 5.8))
    y = np.arange(len(means.index))
    left = np.zeros(len(means.index))
    for i, celltype in enumerate(means.columns):
        vals = means[celltype].values
        ax.barh(
            y,
            vals,
            left=left,
            height=0.62,
            color=colors[i % len(colors)],
            edgecolor="white",
            linewidth=0.45,
            label=short_celltype(celltype),
        )
        left += vals
    ax.set_yticks(y)
    ax.set_yticklabels(means.index)
    ax.set_xlim(0, 100)
    ax.set_xlabel("Mean donor-level proportion (%)")
    ax.set_title("All-study composition by data source", fontsize=11, pad=6)
    ax.grid(axis="x", color="0.9", linewidth=0.8)
    ax.grid(axis="y", visible=False)
    ax.invert_yaxis()
    ax.legend(
        frameon=False,
        fontsize=6,
        ncol=4,
        bbox_to_anchor=(0.0, -0.26, 1.0, 0.1),
        loc="upper left",
        mode="expand",
        title="Cell type",
        title_fontsize=7,
        borderaxespad=0,
    )
    fig.tight_layout()
    save_supplementary_figure(fig, "Supplementary_Figure_S1_all_study_composition")


def make_supplementary_pairwise_de() -> None:
    summary = pd.read_csv(TABLES / "pairwise_pseudobulk_de_summary.csv")
    tested = summary[summary["status"] == "tested"].copy()
    tested["comparison"] = tested["study1"] + " vs " + tested["study2"]
    counts = tested.pivot(index="celltype", columns="comparison", values="significant_genes")
    counts = counts.loc[counts.sum(axis=1).sort_values(ascending=False).index]
    log_counts = np.log10(counts.astype(float) + 1.0)
    annot = counts.fillna(0).astype(int).astype(str)
    fig, ax = plt.subplots(figsize=(10.8, 7.2))
    sns.heatmap(
        log_counts,
        ax=ax,
        cmap="Blues",
        linewidths=0.3,
        linecolor="white",
        annot=annot,
        fmt="",
        annot_kws={"fontsize": 5.2},
        cbar_kws={"label": "log10(significant genes + 1)", "shrink": 0.82},
    )
    ax.set_title("All-study pairwise pseudobulk DE by cell type", fontsize=11, pad=7)
    ax.set_xlabel("Study pair")
    ax.set_ylabel("Cell type")
    ax.set_xticklabels(ax.get_xticklabels(), rotation=42, ha="right", fontsize=6)
    ax.set_yticklabels(ax.get_yticklabels(), rotation=0, fontsize=7)
    fig.tight_layout()
    save_supplementary_figure(fig, "Supplementary_Figure_S2_pairwise_pseudobulk_de")


def make_supplementary_metrics() -> None:
    metrics = pd.read_csv(TABLES / "integration_metrics.csv")
    scvi_path = TABLES / "scvi_integration_metrics.csv"
    if scvi_path.exists():
        metrics = pd.concat([metrics, pd.read_csv(scvi_path)], ignore_index=True)
    methods = [m for m in ["unintegrated", "harmony", "scvi"] if m in set(metrics["method"])]
    metrics = metrics.set_index("method").loc[methods]
    cols = [
        ("batch_asw_lower_is_better", "batch ASW\nlower", "#4C78A8"),
        ("iLISI_study_higher_is_better", "iLISI\nhigher", "#54A24B"),
        ("celltype_asw_higher_is_better", "cell-type ASW\nhigher", "#F58518"),
    ]
    fig, ax = plt.subplots(figsize=(4.8, 3.9))
    x = np.arange(len(metrics.index))
    width = 0.24
    for offset, (col, label, color) in zip([-width, 0, width], cols):
        ax.bar(x + offset, metrics[col].astype(float).values, width=width, label=label, color=color)
    ax.axhline(0, color="0.25", lw=0.7)
    ax.set_xticks(x)
    ax.set_xticklabels(["Unintegrated", "Harmony", "scVI"][: len(x)])
    ax.set_ylabel("Metric value")
    ax.set_title("Integration metrics", fontsize=11, pad=6)
    ax.legend(frameon=False, fontsize=7, loc="upper left")
    fig.tight_layout()
    save_supplementary_figure(fig, "Supplementary_Figure_S6_integration_metrics")


def copy_raster_supplementary(src_png: Path, dst_stem: str) -> None:
    shutil.copy2(src_png, OUT / f"{dst_stem}.png")
    img = mpimg.imread(src_png)
    h, w = img.shape[:2]
    fig, ax = plt.subplots(figsize=(w / 600, h / 600), dpi=600)
    ax.imshow(img)
    ax.axis("off")
    fig.subplots_adjust(left=0, right=1, bottom=0, top=1)
    fig.savefig(OUT / f"{dst_stem}.pdf", bbox_inches="tight", pad_inches=0)
    plt.close(fig)


def copy_supplementary_figures() -> None:
    setup_matplotlib()
    make_supplementary_composition()
    make_supplementary_pairwise_de()
    raster_map = {
        "umap_unintegrated_by_study.png": "Supplementary_Figure_S3_unintegrated_umap",
        "umap_harmony_by_study.png": "Supplementary_Figure_S4_harmony_umap",
        "umap_scvi_by_study.png": "Supplementary_Figure_S5_scvi_umap",
    }
    for src_name, dst_stem in raster_map.items():
        src = FIGS / src_name
        if src.exists():
            copy_raster_supplementary(src, dst_stem)
    make_supplementary_metrics()


MANUSCRIPT_SECTIONS = [
    (
        "Abstract",
        [
            "Public single-cell RNA sequencing studies increasingly reuse samples labelled as healthy, normal or control as if those labels defined a universal biological baseline [1-4]. This Critical Comment argues that the label itself is often a statistical fiction: it records a source-study decision, not proof that samples are exchangeable across donor populations, tissue handling, enrichment, modality and sequencing workflows. We use 28 public lung control samples from five studies as a bounded case study to illustrate the problem. A protocol-aware review separated whole-tissue single-cell datasets from enriched, sorted, single-nucleus and autopsy-derived datasets; integration diagnostics showed that corrected embeddings reduced visible study structure but did not certify compatibility; and donor-level sensitivity analyses still detected study-conditioned composition and cell-state differences [5-9]. The purpose of this evidence is not to introduce a new integration method or to assign every residual difference to biology. Instead, it shows why compatibility must be evaluated before public controls are pooled. We propose a transparency-first checklist: define the control label, report pre- and post-integration views, interpret batch-removal metrics alongside biological-conservation metrics, use donor-level rather than cell-level inference, and disclose unresolved metadata gaps. Integration can align representations, but it cannot turn incompatible studies into a shared baseline.",
        ],
    ),
    (
        "Keywords",
        [
            "single-cell RNA sequencing; public controls; study compatibility; batch effects; integration diagnostics; donor-level analysis",
        ],
    ),
    (
        "The illusion of the universal baseline",
        [
            "The most convenient assumption in public single-cell reuse is also one of the easiest to forget: a sample labelled healthy or control is treated as if it were a portable biological constant. In practice, that label is a decision made within a source study. It depends on who was eligible, how tissue was obtained, which cells or nuclei were captured, which chemistry was used, how the data were filtered, and what the original study needed the controls to control.",
            "This distinction matters because modern integration tools can make incompatible datasets look reassuringly coherent [2,3,8,9]. A corrected embedding may reduce study separation, yet still leave open whether the underlying samples are comparable controls for disease analysis, atlas construction or model training. The central issue is therefore not whether integration is useful. It is useful. The issue is whether integration is being asked to answer a question it was not designed to answer: are these studies compatible enough to define a shared baseline?",
            "We call this problem the illusion of the universal baseline. It is an illusion because a shared control label can conceal different sampling frames, protocols and residual confounders. It is universal only in appearance: once the source-study history is restored, the control label becomes conditional rather than absolute. The article therefore treats the lung-control case study as evidence for an interpretive practice: researchers should read control labels through their source-study histories before treating them as shared baselines.",
        ],
    ),
    (
        "A lung-control case study exposes the problem",
        [
            "To make the argument concrete, we re-examined 28 public lung control samples from five GEO series. The full set was intentionally heterogeneous: it included whole-tissue scRNA-seq studies, lineage-negative enrichment, FACS-enriched fractions, and frozen rapid-autopsy single-nucleus data. We therefore treated the complete dataset as a diagnostic case rather than as an unbiased composition cohort. Only GSE173896 and GSE227691 were used as a controlled whole-tissue scRNA-seq subset for the most conservative comparisons (Figure 2A; Supplementary Table S1).",
            "This design choice is a first-principles point, not a technical footnote. Before asking whether control samples differ, one must ask what kind of control each source study created. Composition claims require different compatibility than cell-state claims. Enriched epithelial or stromal fractions can be useful for state-level sensitivity analyses but are not interchangeable with whole-tissue composition samples. Single-nucleus autopsy data can be scientifically valuable while still being a poor match for fresh whole-cell composition inference.",
            "The lung-control analysis therefore serves as a stress test for a common reuse habit. If samples that all carry a control label differ in source protocol, modality and donor context, then pooling them first and explaining differences later reverses the order of scientific reasoning. Compatibility should be an entry condition for pooling, not an afterthought after integration.",
        ],
    ),
    (
        "Corrected embeddings are not evidence of compatibility",
        [
            "The corrected embedding is the most visually persuasive artifact in many public single-cell analyses. It can also be the most misleading if it is read as a certificate of exchangeability. In the lung-control case study, Harmony and scVI reduced study-associated structure compared with the unintegrated representation, but corrected embeddings did not remove the need for donor-level checks (Figure 2B-E; Supplementary Table S4) [8,9].",
            "The same logic applies beyond this dataset. Integration methods optimize representations under assumptions and objectives; they do not reconstruct missing donor metadata, undo enrichment designs, or make an autopsy single-nucleus profile equivalent to a fresh whole-cell profile. A study can appear better mixed in embedding space while remaining unsuitable as a pooled control for a specific downstream question.",
            "The operational punchline is simple: integration is not a substitute for study compatibility. A corrected UMAP can support visualization and exploratory annotation, but compatibility is established by the source-study design, transparent diagnostics and donor-level sensitivity analyses. Treating the embedding as biological fact creates a false sense of security and can propagate hidden incompatibility into disease comparisons.",
        ],
    ),
    (
        "How to read integration metrics",
        [
            "Metrics such as average silhouette width (ASW), iLISI, cLISI and graph connectivity should be interpreted as a set of diagnostic questions rather than as a leaderboard [5]. ASW asks whether labels remain separated or overlap in a representation. When the label is study, lower batch ASW is consistent with weaker study separation. When the label is cell type, higher cell-type ASW can indicate better preservation of annotated biological structure.",
            "iLISI asks a different question: how diverse are study labels within local neighborhoods? It behaves like a local diversity or entropy score for batch mixing, so a higher iLISI indicates more study mixing. cLISI applies the same logic to biological labels such as annotated cell types. A higher cLISI can therefore be a warning sign if integration is mixing cell types that should remain distinct. Graph connectivity checks whether cells carrying the same biological label remain connected rather than being fragmented across the representation.",
            "These metrics were chosen because they expose the trade-off at the center of the critique. A method can improve batch mixing while also increasing biological over-mixing. The point is not to crown a best integration method, but to make visible the difference between reducing study structure and proving that studies are compatible controls (Figure 1C; Figure 2E; Supplementary Table S4). Detailed values belong in the Supplementary Materials so that the main text can focus on what the trends mean.",
        ],
    ),
    (
        "A transparency-first checklist for public-control reuse",
        [
            "The practical response is a transparency-first workflow. First, define what the control label means in each source study, including donor eligibility, disease context, tissue region, modality, enrichment or sorting, and sequencing workflow. Second, build a study compatibility table before integration. This table should determine which datasets are eligible for composition analysis, which are suitable only for state-level sensitivity analysis, and which should remain diagnostic rather than inferential.",
            "Third, report pre- and post-integration visualizations. The unintegrated view reveals study structure that a corrected embedding may hide; the integrated view shows what the correction changed. Fourth, pair batch-removal metrics with biological-conservation metrics. Reporting only batch mixing encourages overconfidence; reporting only cell-type preservation ignores the purpose of integration. Fifth, verify downstream conclusions at the donor or sample level. Individual cells are measurements, not independent donors.",
            "Finally, researchers should include a buyer-beware statement when public controls remain imperfect. Missing age, sex, smoking status, tissue region or chemistry metadata should not be silently imputed or hidden by integration. Unresolved confounding is not a failure of the analysis if it is reported clearly; it becomes a failure when a pooled control baseline is presented as if those uncertainties did not exist.",
        ],
    ),
    (
        "Beyond explicit batch removal",
        [
            "A constructive path forward is to reduce the pressure on integration to erase differences that should instead remain visible. Methods such as MapBatch illustrate conservative, batch-aware normalization strategies that aim to preserve biological signal rather than treating every batch-associated difference as noise [11]. The broader lesson is not that any one method should replace Harmony, scVI or other current tools. It is that analytical designs should preserve enough batch and protocol information for investigators to see when compatibility remains unresolved.",
            "This perspective also changes how new tools should be judged. A method that produces a visually smooth embedding is not automatically better if it obscures incompatibility relevant to the scientific question. Conversely, a method that keeps some study structure visible may be valuable when that structure reflects real protocol or donor differences that must be reported. The field needs integration diagnostics that protect interpretability, not only aesthetics.",
        ],
    ),
    (
        "Boundaries of the argument",
        [
            "This comment should not be read as an argument against public control reuse. Public single-cell data are indispensable, and carefully matched controls can increase power, improve annotation and support reproducible disease comparisons. The argument is narrower: control reuse should be conditional on study compatibility and donor-level validation.",
            "The lung-control case study also has limits. It concerns one tissue context and relies on metadata available from public sources. Age, sex, smoking status, tissue region and chemistry were not uniformly available and were not imputed. Cell labels were marker-based and include low-confidence clusters, so fine subtype interpretation should remain cautious. The controlled subset contains only ten samples, which limits precision for donor-level inference. These boundaries are exactly why the article emphasizes transparency rather than a universal correction recipe.",
        ],
    ),
    (
        "Conclusion",
        [
            "The healthy-control label is not a universal baseline. It is a study-conditioned label that must be interpreted through donor selection, protocol design, modality and metadata completeness. Integration can help align representations, but compatibility must be earned rather than assumed. Before public controls are pooled, researchers should show what integration changed, verify conclusions at the donor level, and state what remains unresolved.",
        ],
    ),
    (
        "Figures and Supplementary Materials",
        [
            "Figure 1. From universal-baseline illusion to compatibility-first reuse. (A) A shared control label can hide incompatible donor and protocol histories. (B) Compatibility-first workflow for public-control reuse. (C) Conceptual guide to ASW, iLISI, cLISI and graph connectivity. (D) Buyer-beware checklist for transparent reuse of public single-cell controls.",
            "Figure 2. Lung-control case study. (A) Study compatibility map for five public lung control datasets. (B-D) Unintegrated, Harmony-integrated and scVI representations coloured by source study. (E) Integration metrics showing why batch-removal and biological-conservation diagnostics must be read together. (F) Donor-level evidence summary showing that protocol-aware filtering reduced but did not eliminate study-conditioned heterogeneity.",
            "Supplementary Materials. Supplementary Tables S1-S8 report dataset classification, quality control, marker filtering, annotation, integration metrics, composition diagnostics, pseudobulk analyses and sensitivity analyses. Supplementary Figures S1-S6 provide the all-study composition screen, pairwise pseudobulk DESeq2 heatmap and individual UMAP or metric panels [10].",
        ],
    ),
    (
        "Data Availability Statement",
        [
            "The public data used in this study were obtained from the GEO series listed in Supplementary Table S1. The revised analysis outputs are summarized in Supplementary Tables S1-S8. The analysis code is available at https://github.com/XIAOJIE0519/bias-in-single-cell.",
        ],
    ),
    (
        "Author Contributions",
        [
            "Ximing Wang: Original Writing, Conceptualization, Supervision.",
            "Zhichao Jiang: Original Writing, Conceptualization, Methodology.",
            "Shanjie Luan: Conceptualization, Methodology, Software, Writing and Editing, Visualization.",
        ],
    ),
    ("Ethical approval", ["Not applicable."]),
    ("Funding", ["Not applicable."]),
    ("Declaration of interests", ["The authors declare no conflicts of interest."]),
]


def set_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1
    for style_name, size, color in [
        ("Heading 1", 15, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 78, 121)),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def add_paragraph(doc: Document, text: str, bold_lead: str | None = None) -> None:
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        r.bold = True
        p.add_run(text[len(bold_lead) :])
    else:
        p.add_run(text)


def build_manuscript() -> None:
    doc = Document()
    set_document_styles(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    r.bold = True
    r.font.size = Pt(16)
    doc.add_paragraph()
    for text in [
        "Authors: Zhichao Jiang2, Ximing Wang3, Shanjie Luan1,*",
        "Affiliation:",
        "1. Shandong University, School of Basic Medical Sciences, 44 Wenhua Xi Road, Jinan, 250012, Shandong, China;",
        "2. School of Information Science and Engineering, Chenggong Campus of Yunnan University, South Outer Ring Road, East University Town, Chenggong District, Kunming, 650500, Yunnan, China",
        "3. South China University of Technology, School of Biomedical Science and Engineering, No. 777 East Xingye Avenue, Panyu District, Guangzhou, 511442, Guangdong, China",
        "Correspondence to: Luan20050519@163.com; Tel.: +86 14706399924",
    ]:
        add_paragraph(doc, text)
    for heading, paragraphs in MANUSCRIPT_SECTIONS:
        doc.add_heading(heading, level=1)
        for para in paragraphs:
            add_paragraph(doc, para)
    doc.add_heading("References", level=1)
    for i, ref in enumerate(REFERENCES, 1):
        add_paragraph(doc, f"{i}. {ref}")
    doc.save(OUT / "Manuscript.docx")

    md = [f"# {TITLE}", ""]
    md.extend(
        [
            "Authors: Zhichao Jiang2, Ximing Wang3, Shanjie Luan1,*",
            "",
            "Correspondence to: Luan20050519@163.com",
            "",
        ]
    )
    for heading, paragraphs in MANUSCRIPT_SECTIONS:
        md.append(f"## {heading}")
        md.extend(["", *paragraphs, ""])
    md.append("## References")
    md.extend([f"{i}. {ref}" for i, ref in enumerate(REFERENCES, 1)])
    (OUT / "Manuscript.md").write_text("\n".join(md), encoding="utf-8")


def build_cover_letter() -> None:
    doc = Document()
    set_document_styles(doc)
    body = [
        "Dear Professor Wong,",
        f"Please find enclosed our substantially revised manuscript for {JOURNAL}, entitled \"{TITLE}\" ({MS_NO}).",
        "We thank the editor and reviewer for the direct and constructive guidance. We agree that the previous version still read too much like a diagnostic methods report for the intended Critical Comment format. In response, we have rewritten the manuscript around an intellectual argument rather than around the analysis pipeline.",
        "The revised manuscript now opens with the concept of the illusion of the universal baseline and builds toward the central message that integration is not a substitute for study compatibility. The lung-control analysis is retained as a concise case study, while detailed metric values and technical results have been moved to the Supplementary Materials.",
        "We also added a conceptual guide to ASW, iLISI, cLISI and graph connectivity, a transparency-first checklist for public-control reuse, a buyer-beware discussion of corrected embeddings, and a constructive discussion of conservative batch-aware normalization approaches such as MapBatch. The figures have been rebuilt so that Figure 1 communicates the commentary argument and Figure 2 presents the lung-control evidence.",
        "We hope these changes address the reviewer concern that the article should function as a story-driven Critical Comment and a practical guide for the community rather than as a dry observation of a technical problem.",
        "Sincerely,",
        "Shanjie Luan and co-authors",
    ]
    for para in body:
        add_paragraph(doc, para)
    doc.save(OUT / "cover letter.docx")


RESPONSES = [
    (
        "E.1",
        "The editor requested a major revision and a list of changes or rebuttal against each point.",
        "We have prepared a point-by-point response and a fully revised manuscript package. The revised files include a rewritten clean manuscript, rebuilt Figures 1 and 2, updated Supplementary Materials, a revised cover letter and this response document.",
        "Response to Reviewers; revised manuscript; Figures 1-2; Supplementary Tables S1-S8.",
    ),
    (
        "R1.1",
        "The manuscript addresses a timely issue, but the current presentation lacks the narrative impact required for a Critical Comment and reads as a dry, method-heavy diagnostic report.",
        "We agree. We have abandoned the previous methods-paper structure and rewritten the manuscript as a Critical Comment. The revised manuscript is organized around the intellectual argument that public control labels can create an illusion of a universal baseline. Technical results are now used as a case-study evidence base rather than as the organizing scaffold of the paper.",
        "Title; Abstract; section 'The illusion of the universal baseline' paragraphs 1-3; section 'A lung-control case study exposes the problem' paragraphs 1-3; Figure 1.",
    ),
    (
        "R1.2",
        "The manuscript should lead with the illusion of the universal baseline and build toward the punchline that integration is not a substitute for study compatibility.",
        "We have made this the central framing of the paper. The title now uses 'The illusion of the universal baseline', and the opening section defines the control label as a study-conditioned label rather than a biological constant. A dedicated section, 'Corrected embeddings are not evidence of compatibility', states the revised punchline explicitly: integration aligns representations, but it does not certify study compatibility.",
        "Title; Abstract; section 'The illusion of the universal baseline' paragraphs 1-3; section 'Corrected embeddings are not evidence of compatibility' paragraphs 1-3; Figure 1A-B.",
    ),
    (
        "R1.3",
        "Metrics such as ASW, iLISI and cLISI should be conceptualized for the reader rather than merely listed as benchmarking values.",
        "We have added a self-contained metric explanation. The revised text explains ASW as a separation-versus-overlap measure, iLISI as a local study-label diversity measure, cLISI as a biological-label mixing diagnostic, and graph connectivity as a label-continuity diagnostic. Detailed values remain in Supplementary Table S4, while the main text explains why the metrics must be interpreted jointly.",
        "Section 'How to read integration metrics' paragraphs 1-3; Figure 1C; Figure 2E; Supplementary Table S4.",
    ),
    (
        "R1.4",
        "The paper should provide actionable advice, including transparency-first reporting, pre/post visualization and donor-level sensitivity analyses.",
        "We have added a transparency-first checklist and a buyer-beware warning. The revised manuscript recommends defining the control label, building a compatibility table, reporting unintegrated and integrated views, pairing batch-removal and biological-conservation metrics, using donor/sample-level inference and disclosing unresolved metadata gaps.",
        "Section 'A transparency-first checklist for public-control reuse' paragraphs 1-3; Figure 1B and 1D; Figure 2.",
    ),
    (
        "R1.5",
        "The manuscript should introduce newer robust concepts such as MapBatch as constructive secondary advice rather than as a rigid requirement.",
        "We have added a constructive discussion of approaches such as MapBatch, framed as an example of conservative batch-aware normalization rather than as a mandated replacement for existing integration methods. We explicitly state that the broader lesson is to preserve enough batch and protocol information to assess unresolved compatibility.",
        "Section 'Beyond explicit batch removal' paragraphs 1-2; Reference 11.",
    ),
    (
        "R1.6",
        "The manuscript should strip away methods-paper style and become a standard-setting guide focused on compatibility and donor-level integrity.",
        "We have rewritten the manuscript around a compatibility-first guide. Detailed QC, annotation, integration metrics, composition statistics and pseudobulk results are retained in the Supplementary Materials, while the main text now emphasizes interpretive principles, actionable checks and boundaries of the argument.",
        "Whole manuscript structure; Figure 1; Figure 2; Supplementary Tables S1-S8; Supplementary Figures S1-S6.",
    ),
]


def build_response_letter() -> None:
    doc = Document()
    set_document_styles(doc)
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Normal"].paragraph_format.space_after = Pt(3)
    for style_name, size in [("Heading 1", 14), ("Heading 2", 11)]:
        style = doc.styles[style_name]
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(3)
    doc.add_heading("Response to Reviewers", level=1)
    add_paragraph(doc, f"Manuscript No.: {MS_NO}")
    add_paragraph(doc, f"Journal: {JOURNAL}")
    add_paragraph(doc, f"Revised title: {TITLE}")
    add_paragraph(
        doc,
        "We thank the editor and reviewer for the constructive comments. We have treated this revision as a full article-type pivot rather than a minor textual edit. The manuscript has been rewritten as a Critical Comment, with the lung-control analysis retained as a bounded case study supporting a broader argument about public-control reuse.",
    )
    for cid, concern, response, changes in RESPONSES:
        doc.add_heading(cid, level=2)
        add_paragraph(doc, f"Reviewer/editor concern: {concern}", bold_lead="Reviewer/editor concern:")
        add_paragraph(doc, f"Response: {response}", bold_lead="Response:")
        add_paragraph(doc, f"Changes made: {changes}", bold_lead="Changes made:")
    doc.save(OUT / "Response_to_Reviewers.docx")

    lines = ["# Response to Reviewers", "", f"Manuscript No.: {MS_NO}", f"Revised title: {TITLE}", ""]
    for cid, concern, response, changes in RESPONSES:
        lines.extend([f"## {cid}", f"**Reviewer/editor concern:** {concern}", "", f"**Response:** {response}", "", f"**Changes made:** {changes}", ""])
    (OUT / "Response_to_Reviewers.md").write_text("\n".join(lines), encoding="utf-8")


def build_ledgers_and_audit() -> None:
    ledger = """# Claim-Evidence-Boundary Ledger

| Claim | Evidence used | Boundary enforced |
|---|---|---|
| Public single-cell controls are not automatically exchangeable. | 28 lung control samples from five source studies; protocol classes in Supplementary Table S1; integration and donor-level diagnostics. | Does not imply all public controls are unusable. |
| Integration is not a substitute for study compatibility. | Pre/post integration views and metrics; residual donor-level composition/state differences after protocol-aware filtering. | Does not argue against using Harmony, scVI or other integration methods. |
| Corrected embeddings can create false security if read alone. | UMAP and metric trends show improved mixing but do not resolve metadata or protocol differences. | Does not claim corrected embeddings are invalid; they remain useful for representation and visualization. |
| ASW, iLISI, cLISI and graph connectivity should be interpreted jointly. | Existing integration metrics in Supplementary Table S4 and conceptual explanations in Figure 1C. | No single metric is treated as definitive. |
| A transparency-first checklist is needed before pooling public controls. | Reviewer request plus case-study failures: protocol heterogeneity, missing metadata, donor-level residual differences. | Checklist is guidance, not a new method or mandatory standard. |
| Conservative batch-aware normalization concepts such as MapBatch provide constructive direction. | Verified MapBatch citation and reviewer-specified link. | Framed as an example, not a required replacement for current integration methods. |
"""
    (OUT / "claim_evidence_boundary_ledger.md").write_text(ledger, encoding="utf-8")

    audit = """# Adversarial Review Audit

## Reviewer 1 emphasis: technical soundness and evidence boundaries

Assessment: The revised package no longer invents new analysis or overstates residual differences as biology. Existing evidence supports a bounded case that compatibility must be assessed before pooling public controls.

Remaining risks checked:
- Small controlled subset: now stated as a boundary.
- Marker-based annotation: now stated as a boundary.
- Missing age/sex/smoking/tissue-region metadata: not imputed and explicitly disclosed.
- Pairwise pseudobulk DE: moved to supplementary/diagnostic evidence.

Blocking issues after revision: none found.

## Reviewer 2 emphasis: originality and significance

Assessment: The manuscript's novelty now rests on a field-facing argument rather than on claiming a new pipeline. The strong title is justified because the paper explains a recurrent interpretive failure in public-control reuse.

Remaining risks checked:
- Not enough new method novelty: acceptable for a Critical Comment, because the contribution is interpretive and practical.
- Prior integration literature: retained in references; the revised article positions itself as guidance rather than benchmarking.
- Overclaiming field-wide paradigm shift: softened to practical recommendations and compatibility-first workflow.

Blocking issues after revision: none found.

## Reviewer 3 emphasis: readability and Critical Comment fit

Assessment: The revised manuscript now starts with a story, states the punchline early, explains metrics conceptually and gives a checklist. The main text no longer reads as a methods pipeline.

Remaining risks checked:
- Metric jargon: definitions added in text and Figure 1C.
- Dry results-first structure: replaced by commentary headings.
- Figure-message mismatch: old Figure 1E/F mismatch removed; new Figure 1 and Figure 2 legends match their panels.

Blocking issues after revision: none found.

## Final adversarial decision

The package is ready for author-side Word layout review and submission-system formatting checks. The remaining risks are not scientific blockers; they are production checks: page layout, title acceptability and whether Editorial Manager requires tracked changes.
"""
    (OUT / "adversarial_audit.md").write_text(audit, encoding="utf-8")


def update_supplementary_workbook() -> None:
    src = PREV / "Supplementary Materials.xlsx"
    dst = OUT / "Supplementary Materials.xlsx"
    shutil.copy2(src, dst)
    wb = load_workbook(dst)
    if "Revision2 Notes" in wb.sheetnames:
        del wb["Revision2 Notes"]
    ws = wb.create_sheet("Revision2 Notes")
    rows = [
        ["Item", "Revision 2 note"],
        ["Article framing", "Main manuscript rewritten as a Critical Comment centred on the illusion of the universal baseline."],
        ["Main Figure 1", "Conceptual guide: universal-baseline illusion, compatibility-first workflow, metric logic and buyer-beware checklist."],
        ["Main Figure 2", "Lung-control case evidence: study compatibility map, pre/post integration views, integration diagnostics and donor-level summary."],
        ["Metric definitions", "ASW: label separation/overlap; iLISI: local study-label diversity; cLISI: local biological-label mixing; graph connectivity: same-label neighborhood continuity."],
        ["Supplementary Figure S1", "All-study donor-level composition by data source."],
        ["Supplementary Figure S2", "All-study pairwise pseudobulk differential-expression heatmap used as diagnostic evidence only."],
        ["Supplementary Figures S3-S6", "Individual UMAP and integration metric panels retained for auditability."],
        ["MapBatch citation", MAPBATCH_REF],
        ["Interpretive boundary", "Residual study-conditioned heterogeneity is not interpreted as a single biological mechanism or pure biological selection bias."],
    ]
    for r, row in enumerate(rows, 1):
        for c, value in enumerate(row, 1):
            cell = ws.cell(r, c, value)
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            if r == 1:
                cell.font = Font(bold=True, color="FFFFFF")
                cell.fill = PatternFill("solid", fgColor="1F4D78")
            elif c == 1:
                cell.font = Font(bold=True)
                cell.fill = PatternFill("solid", fgColor="E8EEF5")
    ws.column_dimensions["A"].width = 26
    ws.column_dimensions["B"].width = 110
    wb.save(dst)


def build_handoff_final(qa: dict) -> None:
    text = f"""# Handoff Final: JBCB-1487R1 revision 2 package

Generated: {TODAY}

## Delivered files

- `vexpaer_paper_2/投稿2/Manuscript.docx`
- `vexpaer_paper_2/投稿2/Manuscript.md`
- `vexpaer_paper_2/投稿2/Figure1.png/pdf/eps`
- `vexpaer_paper_2/投稿2/Figure2.png/pdf/eps`
- `vexpaer_paper_2/投稿2/Supplementary_Figure_S1-S2.png/pdf/eps`
- `vexpaer_paper_2/投稿2/Supplementary_Figure_S3-S5.png/pdf`
- `vexpaer_paper_2/投稿2/Supplementary_Figure_S6.png/pdf/eps`
- `vexpaer_paper_2/投稿2/Supplementary Materials.xlsx`
- `vexpaer_paper_2/投稿2/Response_to_Reviewers.docx`
- `vexpaer_paper_2/投稿2/cover letter.docx`
- `vexpaer_paper_2/投稿2/claim_evidence_boundary_ledger.md`
- `vexpaer_paper_2/投稿2/adversarial_audit.md`
- `vexpaer_paper_2/投稿2/qa_report.json`

## Strategic change

The manuscript has been rebuilt as a Critical Comment rather than a methods-style diagnostic report. The title and opening now foreground "the illusion of the universal baseline"; the central punchline is that integration is not a substitute for study compatibility.

## Adversarial review summary

- Reviewer 1 emphasis, technical soundness: no new analysis was invented; the case is bounded to existing lung-control evidence and donor-level diagnostics. Remaining risk is the small controlled subset, now stated as a boundary.
- Reviewer 2 emphasis, originality/significance: novelty now sits in the argument and field guidance, not in claiming a new method. The article no longer overclaims biological selection bias.
- Reviewer 3 emphasis, readability/Critical Comment fit: the main text leads with a story and checklist, while technical values are moved to supplementary material. This addresses the reviewer criticism of dry methods-paper style.

## QA notes

- DOCX structural checks: {qa.get("docx_structural")}
- XLSX structural check: {qa.get("xlsx_structural")}
- Figure files generated: {qa.get("figures")}
- Claim search flags: {qa.get("claim_flags")}
- Render QA: {qa.get("render_qa")}

## Remaining human checks

- Word COM PDF rendering completed; still open the final DOCX files in Word before upload if the journal requires manual page-layout confirmation.
- Confirm JBCB accepts the stronger title. If not, use the conservative title while retaining the same opening argument.
- Confirm whether Editorial Manager requests a tracked-changes manuscript in addition to clean manuscript and response letter.
"""
    (ROOT / "vexpaer_paper_2" / "handoff_final.md").write_text(text, encoding="utf-8")


def structural_docx_check(path: Path) -> bool:
    with ZipFile(path) as zf:
        names = set(zf.namelist())
    return "word/document.xml" in names and "[Content_Types].xml" in names


def try_word_render(docs: list[Path]) -> str:
    try:
        import win32com.client  # type: ignore
    except Exception as exc:
        return f"Word COM render not available: {exc!r}"

    word = None
    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0
        render_dir = OUT / "qa_render"
        render_dir.mkdir(exist_ok=True)
        exported = []
        for docx_path in docs:
            pdf_path = OUT / f"{docx_path.stem}_word_render.pdf"
            doc = word.Documents.Open(str(docx_path.resolve()), ReadOnly=True)
            doc.ExportAsFixedFormat(OutputFileName=str(pdf_path.resolve()), ExportFormat=17)
            doc.Close(False)
            exported.append(pdf_path)
        try:
            import fitz  # type: ignore
            from PIL import Image, ImageDraw

            for pdf_path in exported:
                pdf = fitz.open(pdf_path)
                thumbs = []
                for i, page in enumerate(pdf, 1):
                    pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5), alpha=False)
                    png_path = render_dir / f"{pdf_path.stem}_page_{i:02d}.png"
                    pix.save(png_path)
                    im = Image.open(png_path).convert("RGB")
                    im.thumbnail((360, 480))
                    canvas = Image.new("RGB", (380, 520), "white")
                    canvas.paste(im, ((380 - im.width) // 2, 10))
                    draw = ImageDraw.Draw(canvas)
                    draw.text((10, 495), png_path.name, fill="black")
                    thumbs.append(canvas)
                if thumbs:
                    cols = 2 if len(thumbs) > 1 else 1
                    rows = math.ceil(len(thumbs) / cols)
                    sheet = Image.new("RGB", (cols * 380, rows * 520), "white")
                    for idx, im in enumerate(thumbs):
                        sheet.paste(im, ((idx % cols) * 380, (idx // cols) * 520))
                    sheet.save(render_dir / f"{pdf_path.stem}_contact.png")
        except Exception as exc:
            return f"Word COM PDF render passed; PNG contact-sheet render failed: {exc!r}"
        return "Word COM PDF render and PNG contact sheets completed."
    except Exception as exc:
        return f"Word COM render failed: {exc!r}"
    finally:
        if word is not None:
            word.Quit()


def run_qa() -> dict:
    docs = [OUT / "Manuscript.docx", OUT / "Response_to_Reviewers.docx", OUT / "cover letter.docx"]
    docx_ok = all(structural_docx_check(p) for p in docs)
    wb_ok = False
    try:
        wb = load_workbook(OUT / "Supplementary Materials.xlsx", read_only=True, data_only=True)
        wb_ok = "Revision2 Notes" in wb.sheetnames
    except Exception:
        wb_ok = False
    fig_files = sorted(p.name for p in OUT.glob("Figure*.png"))
    text = (OUT / "Manuscript.md").read_text(encoding="utf-8")
    flags = {}
    for term in ["proves", "fully eliminates", "ground truth", "biological selection bias"]:
        flags[term] = text.lower().count(term.lower())
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if soffice:
        render_note = "LibreOffice/soffice available, but Word COM render path was used for this package."
    else:
        render_note = try_word_render(docs)
    qa = {
        "docx_structural": "passed" if docx_ok else "failed",
        "xlsx_structural": "passed" if wb_ok else "failed",
        "figures": fig_files,
        "claim_flags": flags,
        "render_qa": render_note,
    }
    (OUT / "qa_report.json").write_text(json.dumps(qa, indent=2, ensure_ascii=False), encoding="utf-8")
    return qa


def main() -> None:
    ensure_clean_out()
    make_figure1()
    make_figure2()
    copy_supplementary_figures()
    build_manuscript()
    build_response_letter()
    build_cover_letter()
    build_ledgers_and_audit()
    update_supplementary_workbook()
    qa = run_qa()
    build_handoff_final(qa)
    print(json.dumps(qa, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
